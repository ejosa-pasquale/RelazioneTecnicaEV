\
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Dict, List, Optional, Union, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


# -------------------- Data models --------------------
@dataclass
class ProgettistaData:
    nome: str
    indirizzo: str
    cell: str
    email: str
    piva: str


@dataclass
class EsecutriceData:
    nome: str
    indirizzo: str
    piva: str


@dataclass
class ColonninaItem:
    descrizione: str
    quantita: int = 1


@dataclass
class PhotoItem:
    filename: str
    content: bytes
    caption: str = ""


@dataclass
class AllegatoItem:
    filename: str
    content: bytes
    kind: str  # "pdf" | "docx"


@dataclass
class RelazioneData:
    luogo_data: str
    committente: str
    sito_indirizzo: str
    sito_cap_citta: str
    oggetto: str

    distanza_m: int = 60
    potenza_impegnata_kw: float = 4.0
    ik_trifase_ka: float = 10.0
    ik_monofase_ka: float = 6.0
    cavo_lunghezza_m: int = 60
    cavo_tipo: str = "FG16OM16 3G6 0.6/1kV"

    layout_incluso: str = ""
    layout_escluso: str = ""


# -------------------- Anchors / placeholders --------------------
ANCHORS = {
    "DITTA_ESECUTRICE": "{{DITTA_ESECUTRICE}}",
    "LAYOUT": "{{LAYOUT_DESCRITTIVO}}",
    "COLONNINE": "{{COLONNINE}}",
    "FOTO": "{{FOTO_GALLERY}}",
    "DIAGRAMMA": "{{DIAGRAMMA_IMPIANTO}}",
    "ALLEGATI": "{{ALLEGATI_SCHEDA_TECNICA}}",
}

HEADING_HINTS = {
    "DITTA_ESECUTRICE": ["DITTA ESECUTRICE", "IMPRESA ESECUTRICE"],
    "LAYOUT": ["LAYOUT D'IMPIANTO", "LAYOUT D’IMPIANTO"],
    "COLONNINE": ["LOCALIZZAZIONE DELL'IMPIANTO", "LOCALIZZAZIONE DELL’IMPIANTO", "COLONNINA", "WALLBOX"],
    "FOTO": ["DOCUMENTAZIONE FOTOGRAFICA", "FOTOGRAFICA"],
    "DIAGRAMMA": ["SCHEMA", "DIAGRAMMA", "LAYOUT"],
    "ALLEGATI": ["ALLEGATI", "SCHEDA TECNICA", "SCHEDE TECNICHE"],
}

FIELD_PLACEHOLDERS = {
    "{{LUOGO_DATA}}": lambda d: d.luogo_data,
    "{{COMMITTENTE}}": lambda d: d.committente,
    "{{SITO_INDIRIZZO}}": lambda d: d.sito_indirizzo,
    "{{SITO_CAP_CITTA}}": lambda d: d.sito_cap_citta,
    "{{OGGETTO}}": lambda d: d.oggetto,
    "{{DISTANZA_M}}": lambda d: str(d.distanza_m),
    "{{POTENZA_IMPEGNATA_KW}}": lambda d: f"{d.potenza_impegnata_kw:g}",
    "{{IK_TRIFASE_KA}}": lambda d: f"{d.ik_trifase_ka:g}",
    "{{IK_MONOFASE_KA}}": lambda d: f"{d.ik_monofase_ka:g}",
    "{{CAVO_LUNGHEZZA_M}}": lambda d: str(d.cavo_lunghezza_m),
    "{{CAVO_TIPO}}": lambda d: d.cavo_tipo,
}

# fallback su template "statico": sostituisce anche frasi campione del file originale
SAMPLE_TEXT_MAPPING = lambda d: {
    "Busnago, 06/02/2024": d.luogo_data,
    "Nome": d.committente,
    "via della SS. Annunziata 32/A": d.sito_indirizzo,
    "55100 Lucca": d.sito_cap_citta,
    "nuovo impianto di ricarica per veicolo elettrico": d.oggetto,
    "dista circa 60 metri": f"dista circa {d.distanza_m} metri",
    "pari a 10 kA": f"pari a {d.ik_trifase_ka:g} kA",
    "pari a 6 kA": f"pari a {d.ik_monofase_ka:g} kA",
    "4 kW.": f"{d.potenza_impegnata_kw:g} kW.",
    "60 m di cavo FG16OM16 3G6 0.6/1kV": f"{d.cavo_lunghezza_m} m di cavo {d.cavo_tipo}",
}


# -------------------- Iteration across doc --------------------
def iter_table_paragraphs(doc: Document) -> Iterable:
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nt in cell.tables:
                    for nrow in nt.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                yield np

def iter_header_footer_paragraphs(doc: Document) -> Iterable:
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            for p in hf.paragraphs:
                yield p
            for t in hf.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

def iter_all_paragraphs(doc: Document) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for p in iter_table_paragraphs(doc):
        yield p
    for p in iter_header_footer_paragraphs(doc):
        yield p

def _style_name(p) -> str:
    try:
        return (p.style.name or "").lower()
    except Exception:
        return ""

def _is_toc_paragraph(p) -> bool:
    # Word TOC paragraphs often have styles like TOC 1, TOC 2 ...
    s = _style_name(p)
    if s.startswith("toc"):
        return True
    # The whole TOC in your doc appears inside a table; treat table paragraphs as "likely toc" when very short and numeric.
    return False

def _is_in_table(p) -> bool:
    # check ancestors for table-cell tag 'w:tc'
    try:
        parent = p._p.getparent()
        while parent is not None:
            if parent.tag.endswith('}tc'):
                return True
            parent = parent.getparent()
    except Exception:
        pass
    return False

def doc_contains_text(doc: Document, text: str) -> bool:
    needle = text.lower()
    for p in iter_all_paragraphs(doc):
        if needle in ((p.text or "").lower()):
            return True
    return False


# -------------------- Basic mutations --------------------
def _is_heading_like(text: str) -> bool:
    t = (text or "").strip()
    return bool(re.match(r"^\d+(\.|)\s+", t))

def _delete_paragraph(p):
    p._element.getparent().remove(p._element)

def _wipe_paragraph(p):
    for r in p.runs:
        r.text = ""

def _replace_in_paragraph(p, mapping: Dict[str, str]):
    full = "".join(r.text for r in p.runs) if p.runs else (p.text or "")
    new = full
    for old, newv in mapping.items():
        if old:
            new = new.replace(old, newv)
    if new != full:
        if not p.runs:
            p.text = new
        else:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""

def _replace_everywhere(doc: Document, mapping: Dict[str, str]):
    for p in iter_all_paragraphs(doc):
        _replace_in_paragraph(p, mapping)

def _find_first_paragraph_containing(doc: Document, needles: List[str], prefer_body: bool = True):
    needles_l = [n.lower() for n in needles]
    best = None
    for p in iter_all_paragraphs(doc):
        tx = (p.text or "").lower()
        if not any(n in tx for n in needles_l):
            continue
        # skip TOC paragraphs if possible
        if _is_toc_paragraph(p):
            continue
        if prefer_body:
            # Prefer real body paragraphs (not in table, not header/footer)
            if not _is_in_table(p) and p in doc.paragraphs:
                return p
            best = best or p
        else:
            return p
    return best

def ensure_anchors(doc: Document) -> List[str]:
    """
    Inserisce i marker nel CORPO del testo (non nell'indice/TOC).
    """
    created = []
    for key, marker in ANCHORS.items():
        if doc_contains_text(doc, marker):
            continue
        hints = HEADING_HINTS.get(key, [])
        anchor_after = _find_first_paragraph_containing(doc, hints, prefer_body=True) if hints else None
        new_p = doc.add_paragraph(marker)
        if anchor_after is None:
            created.append(marker)
            continue
        anchor_after._p.addnext(new_p._p)
        created.append(marker)
    return created

def build_field_mapping(data: RelazioneData) -> Dict[str, str]:
    return {k: fn(data) for k, fn in FIELD_PLACEHOLDERS.items()}

def _insert_bullets_after(paragraph, doc: Document, lines: List[str]):
    elm = paragraph._p
    for line in lines:
        if not line.strip():
            continue
        bp = doc.add_paragraph(line.strip(), style="List Bullet" if "List Bullet" in doc.styles else None)
        elm.addnext(bp._p)
        elm = bp._p


# -------------------- Section writers --------------------
def _clear_section_after_heading_body(doc: Document, heading_any: List[str]):
    """Cancella paragrafi dopo il titolo NEL BODY fino al prossimo titolo numerato."""
    title = _find_first_paragraph_containing(doc, heading_any, prefer_body=True)
    if not title:
        return None
    body_paras = list(doc.paragraphs)
    if title not in body_paras:
        return title
    idx = body_paras.index(title)
    to_delete = []
    for p2 in body_paras[idx+1:]:
        if _is_heading_like(p2.text):
            break
        to_delete.append(p2)
    for p2 in reversed(to_delete):
        try:
            _delete_paragraph(p2)
        except Exception:
            pass
    return title

def write_layout(doc: Document, data: RelazioneData):
    # prefer explicit anchor in BODY
    marker = ANCHORS["LAYOUT"]
    p = _find_first_paragraph_containing(doc, [marker], prefer_body=True)
    if p:
        _wipe_paragraph(p)
        anchor = p
    else:
        anchor = _clear_section_after_heading_body(doc, ["LAYOUT D'IMPIANTO", "LAYOUT D’IMPIANTO"])
        if not anchor:
            return

    elm = anchor._p
    if data.layout_incluso.strip():
        lbl = doc.add_paragraph("Incluso:")
        if lbl.runs: lbl.runs[0].bold = True
        elm.addnext(lbl._p); elm = lbl._p
        _insert_bullets_after(lbl, doc, data.layout_incluso.splitlines())
    if data.layout_escluso.strip():
        lbl2 = doc.add_paragraph("Escluso:")
        if lbl2.runs: lbl2.runs[0].bold = True
        elm.addnext(lbl2._p); elm = lbl2._p
        _insert_bullets_after(lbl2, doc, data.layout_escluso.splitlines())

def write_colonnine(doc: Document, colonnine: List[ColonninaItem]):
    marker = ANCHORS["COLONNINE"]
    p = _find_first_paragraph_containing(doc, [marker], prefer_body=True)
    if not p:
        return
    _wipe_paragraph(p)
    if not colonnine:
        return
    lines = [f"n. {c.quantita} — {c.descrizione}" for c in colonnine]
    _insert_bullets_after(p, doc, lines)

def write_ditta_esecutrice(doc: Document, esecutrice: EsecutriceData):
    marker = ANCHORS["DITTA_ESECUTRICE"]
    p = _find_first_paragraph_containing(doc, [marker], prefer_body=True)
    if not p:
        return
    _wipe_paragraph(p)
    lines = []
    if esecutrice.nome.strip(): lines.append(esecutrice.nome.strip())
    if esecutrice.indirizzo.strip(): lines.append(esecutrice.indirizzo.strip())
    if esecutrice.piva.strip(): lines.append(f"P.IVA: {esecutrice.piva.strip()}")
    if lines:
        _insert_bullets_after(p, doc, lines)

def write_foto(doc: Document, photos: List[PhotoItem]):
    marker = ANCHORS["FOTO"]
    p = _find_first_paragraph_containing(doc, [marker], prefer_body=True)
    if not p or not photos:
        return
    _wipe_paragraph(p)
    cols = 2
    rows = (len(photos)+cols-1)//cols
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            if idx >= len(photos):
                cell.text = ""
                continue
            item = photos[idx]
            par = cell.paragraphs[0]
            run = par.add_run()
            run.add_picture(io.BytesIO(item.content), width=Inches(3.2))
            if item.caption.strip():
                cap = cell.add_paragraph(item.caption.strip())
                try:
                    cap.style = doc.styles["Caption"]
                except Exception:
                    pass
            idx += 1
    p._p.addnext(table._tbl)

def write_diagramma(doc: Document, diagram_bytes: Optional[bytes]):
    marker = ANCHORS["DIAGRAMMA"]
    p = _find_first_paragraph_containing(doc, [marker], prefer_body=True)
    if not p or not diagram_bytes:
        return
    _wipe_paragraph(p)
    p.add_run().add_picture(io.BytesIO(diagram_bytes), width=Inches(6.5))

def write_allegati(doc: Document, allegati: List[AllegatoItem]):
    marker = ANCHORS["ALLEGATI"]
    p = _find_first_paragraph_containing(doc, [marker], prefer_body=True)
    if not p or not allegati:
        return
    _wipe_paragraph(p)
    lines = [a.filename for a in allegati]
    _insert_bullets_after(p, doc, lines)


# -------------------- Cover writer --------------------
def insert_cover(doc: Document, data: RelazioneData, progettista: ProgettistaData, esecutrice: Optional[EsecutriceData] = None):
    body = doc._body._element

    # --- Fix duplication ---
    # Some templates already include front-matter (cover + index/TOC + revision
    # page). Older versions of this function always *added* a new cover and
    # moved it to the top, leaving the original front-matter in place (result:
    # duplicated beginning of document).
    #
    # Robust strategy: remove everything in the BODY before the first real
    # numbered heading of the report (e.g. "1." or "1 ") or "CAPITOLO 1".
    # This works even when the template uses section breaks instead of page
    # breaks, or when the cover content is in tables/shapes.

    def _is_start_of_main_content(text: str) -> bool:
        t = (text or "").strip().lower()
        if not t:
            return False
        if re.match(r"^1(\.|)\s+", t):
            return True
        if t.startswith("capitolo 1"):
            return True
        # fallback: some templates write "1 Premessa" (no dot)
        if t.startswith("1") and "premessa" in t:
            return True
        return False

    def _remove_front_matter_if_any():
        body_paras = list(doc.paragraphs)
        target_p = None
        for p in body_paras[:250]:
            if _is_start_of_main_content(p.text):
                target_p = p
                break
        if target_p is None:
            return

        body_list = list(body)
        try:
            cut_idx = body_list.index(target_p._p)
        except Exception:
            return

        if cut_idx <= 0:
            return

        for e in list(body_list[:cut_idx]):
            try:
                body.remove(e)
            except Exception:
                pass

    _remove_front_matter_if_any()

    def add_par(text: str, size: int, bold: bool, align: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        return p

    p_title = add_par("RELAZIONE TECNICA", 22, True, "center")
    add_par(data.oggetto.upper(), 14, True, "center")
    add_par("", 11, False, "center")
    add_par(f"Sito: {data.sito_indirizzo} — {data.sito_cap_citta}", 11, False, "center")
    add_par(f"Committente: {data.committente}", 11, False, "center")
    add_par(data.luogo_data, 11, False, "center")
    add_par("", 11, False, "center")

    p_proj = add_par("PROGETTISTA", 12, True, "left")
    elm = p_proj._p
    for line in [
        progettista.nome,
        progettista.indirizzo,
        f"Cell: {progettista.cell}",
        f"Email: {progettista.email}",
        f"P.IVA: {progettista.piva}",
    ]:
        pp = doc.add_paragraph(line)
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elm.addnext(pp._p); elm = pp._p

    if esecutrice and (esecutrice.nome.strip() or esecutrice.indirizzo.strip() or esecutrice.piva.strip()):
        p_ex = doc.add_paragraph()
        p_ex.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p_ex.add_run("DITTA ESECUTRICE"); r.bold = True
        elm.addnext(p_ex._p); elm = p_ex._p
        for line in [
            esecutrice.nome.strip(),
            esecutrice.indirizzo.strip(),
            (f"P.IVA: {esecutrice.piva.strip()}" if esecutrice.piva.strip() else ""),
        ]:
            if not line:
                continue
            pp = doc.add_paragraph(line)
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elm.addnext(pp._p); elm = pp._p

    pb = doc.add_page_break()

    # move cover to top
    body_list = list(body)
    i0 = body_list.index(p_title._p)
    i1 = body_list.index(pb._p)
    elems = body_list[i0:i1+1]
    for e in elems:
        body.remove(e)
    for e in reversed(elems):
        body.insert(0, e)
def prepare_template(template: Union[bytes, Path]) -> bytes:
    doc = Document(io.BytesIO(template) if isinstance(template, (bytes, bytearray)) else str(template))
    ensure_anchors(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def generate_document(
    template: Union[bytes, Path],
    data: RelazioneData,
    progettista: ProgettistaData,
    esecutrice: Optional[EsecutriceData],
    colonnine: List[ColonninaItem],
    photos: List[PhotoItem],
    diagram_bytes: Optional[bytes],
    allegati: List[AllegatoItem],
    extra_fields: Optional[Dict[str, str]] = None,
) -> bytes:
    doc = Document(io.BytesIO(template) if isinstance(template, (bytes, bytearray)) else str(template))

    # 1) anchors in body
    ensure_anchors(doc)

    # 2) replace placeholders if present
    _replace_everywhere(doc, build_field_mapping(data))

    # 3) also replace common static strings (template originale)
    _replace_everywhere(doc, SAMPLE_TEXT_MAPPING(data))

    # 4) cover
    # 2b) sostituzione campi 'aggiungere XXX' (template-driven)
    if extra_fields:
        # sostituisci sia 'aggiungere XXX' che eventuali 'AGGIUNGERE XXX'
        mapping = {}
        for label, value in extra_fields.items():
            token = f"aggiungere {label}".strip()
            mapping[token] = value
            mapping[token.capitalize()] = value
            mapping[token.upper()] = value
        _replace_everywhere(doc, mapping)

    insert_cover(doc, data, progettista, esecutrice)

    # 5) sections (in BODY, not TOC)
    if esecutrice:
        write_ditta_esecutrice(doc, esecutrice)
    write_layout(doc, data)
    write_colonnine(doc, colonnine)
    write_foto(doc, photos)
    write_diagramma(doc, diagram_bytes)
    write_allegati(doc, allegati)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
